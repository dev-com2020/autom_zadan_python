import docx
from datetime import datetime

context = {
    'date': datetime.now(),
    'games': ['Minecraft', 'Cyberpunk', 'Wiedźmin'],
    'total_minutes': 404,
}

document = docx.Document()
document.add_heading('Raport na temat gier',0)
paragraph = document.add_paragraph('Data: ')
paragraph.add_run(str(context['date'])).italic = True
paragraph = document.add_paragraph('Liczba gier w ciągu 30 dni: ')
paragraph.add_run(str(len(context['games']))).italic = True
for game in context['games']:
    document.add_paragraph(game, style='List Bullet')
document.save('raport_word.docx')